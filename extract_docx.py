import sys
from docx import Document

def extract_docx_content(file_path):
    """Extract text content from a .docx file with formatting information."""
    doc = Document(file_path)

    output = []

    # Extract paragraphs with style information
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # Only include non-empty paragraphs
            style = para.style.name if para.style else "Normal"
            output.append(f"[{i}] [{style}] {text}")

    # Also extract tables if present
    if doc.tables:
        output.append("\n\n=== TABLES ===\n")
        for table_idx, table in enumerate(doc.tables):
            output.append(f"\nTable {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                output.append(f"  Row {row_idx + 1}: {' | '.join(cells)}")

    return "\n".join(output)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>")
        sys.exit(1)

    file_path = sys.argv[1]
    content = extract_docx_content(file_path)
    print(content)
