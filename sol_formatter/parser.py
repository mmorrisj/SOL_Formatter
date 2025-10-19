"""Document parser for extracting SOL standards from .docx files"""

from docx import Document
from pathlib import Path
import re
from typing import Dict, List, Any
import json


class SOLParser:
    """Parser for Virginia Standards of Learning documents"""

    def __init__(self):
        self.current_doc = None

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a SOL .docx document and extract structured data

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary containing parsed SOL data
        """
        doc = Document(file_path)
        filename = Path(file_path).name

        # Determine document type from filename
        doc_info = self._parse_filename(filename)

        # Extract content from document
        content = self._extract_content(doc)

        return {
            "metadata": doc_info,
            "content": content,
            "source_file": filename
        }

    def _parse_filename(self, filename: str) -> Dict[str, str]:
        """Extract metadata from filename"""

        # Pattern for numbered SOL documents (e.g., "1-2023-Approved-Math-SOL.docx")
        sol_pattern = r"(\d+)-(.+?)-(\d{4})-Approved-Math-SOL\.docx"
        # Pattern for instructional guides (e.g., "1. Grade 1 Mathematics Instructional Guide.docx")
        guide_pattern = r"(\d+)\.\s+Grade\s+(\d+)\s+Mathematics\s+Instructional\s+Guide\.docx"
        # Pattern for Understanding the Standards (e.g., "12-AFDA-Understanding the Standards.docx")
        understanding_pattern = r"(\d+)-(.+?)-Understanding\s+the\s+Standards\.docx"
        # Pattern for subject instructional guides (e.g., "11. Algebra 2 Mathematics Instructional Guide.docx")
        subject_guide_pattern = r"(\d+)\.\s+(.+?)\s+Mathematics\s+Instructional\s+Guide\.docx"

        if match := re.match(sol_pattern, filename):
            num, subject, year = match.groups()
            return {
                "type": "Approved SOL Standards",
                "number": num,
                "subject": subject,
                "year": year,
                "grade_level": self._get_grade_level(num, subject)
            }
        elif match := re.match(guide_pattern, filename):
            num, grade = match.groups()
            return {
                "type": "Instructional Guide",
                "number": num,
                "grade": grade,
                "subject": "Mathematics",
                "grade_level": f"Grade {grade}"
            }
        elif match := re.match(understanding_pattern, filename):
            num, subject = match.groups()
            return {
                "type": "Understanding the Standards",
                "number": num,
                "subject": subject,
                "grade_level": self._get_grade_level(num, subject)
            }
        elif match := re.match(subject_guide_pattern, filename):
            num, subject = match.groups()
            return {
                "type": "Instructional Guide",
                "number": num,
                "subject": subject,
                "grade_level": self._get_grade_level(num, subject)
            }
        else:
            return {
                "type": "Unknown",
                "filename": filename
            }

    def _get_grade_level(self, num: str, subject: str) -> str:
        """Map document number to grade level or subject"""
        num = int(num)

        if 1 <= num <= 8:
            return f"Grade {num}"

        subject_map = {
            9: "Algebra 1",
            10: "Geometry",
            11: "AFDA",
            12: "Algebra 2",
            13: "Trigonometry",
            14: "Computational Mathematics",
            15: "Probability & Statistics",
            16: "Discrete Mathematics",
            17: "Mathematical Analysis",
            18: "Data Science"
        }

        return subject_map.get(num, subject)

    def _extract_content(self, doc: Document) -> Dict[str, Any]:
        """Extract structured content from document"""

        paragraphs = []
        tables = []
        standards = []

        # Extract all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": para.style.name if para.style else None
                })

                # Try to identify SOL standards (typically formatted like "1.1", "G.5", etc.)
                if self._is_standard(text):
                    standards.append(text)

        # Extract tables (often contain standards and descriptions)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Skip empty rows
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "identified_standards": standards,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables)
        }

    def _is_standard(self, text: str) -> bool:
        """Check if text appears to be a SOL standard identifier"""
        # Common patterns: "1.1", "G.5", "AII.7", "PS.10", etc.
        standard_patterns = [
            r"^[A-Z]{1,4}\.\d+",  # Letter(s) followed by number (e.g., G.5, AII.7)
            r"^\d+\.\d+[a-z]?",   # Number.number with optional letter (e.g., 1.1, 2.3a)
        ]

        for pattern in standard_patterns:
            if re.match(pattern, text):
                return True
        return False

    def save_json(self, data: Dict[str, Any], output_path: str):
        """Save parsed data as JSON"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    def save_csv(self, data: Dict[str, Any], output_path: str):
        """Save parsed standards as CSV"""
        import csv

        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Standard', 'Source', 'Grade Level', 'Subject'])

            for standard in data['content']['identified_standards']:
                writer.writerow([
                    standard,
                    data['source_file'],
                    data['metadata'].get('grade_level', 'Unknown'),
                    data['metadata'].get('subject', 'Mathematics')
                ])
