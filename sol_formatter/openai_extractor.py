"""OpenAI integration for extracting structured data from SOL documents"""

import os
from pathlib import Path
from docx import Document
import json
from typing import Dict, Any, Optional
from openai import OpenAI
from .schema import OPENAI_EXTRACTION_PROMPT, EXTRACTION_SCHEMA

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed, use system env vars only


class OpenAIExtractor:
    """Extract structured data from SOL documents using OpenAI API"""

    def __init__(self, api_key: Optional[str] = None, model: str = "gpt-4o-mini"):
        """
        Initialize OpenAI extractor

        Args:
            api_key: OpenAI API key (defaults to OPENAI_API_KEY env variable)
            model: OpenAI model to use (default: gpt-4o-mini for cost efficiency)
        """
        # Try to load API key from parameter, then .env, then system env
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError(
                "OpenAI API key required. Set OPENAI_API_KEY environment variable "
                "or pass api_key parameter.\n\n"
                "To set up:\n"
                "1. Copy .env.example to .env\n"
                "2. Add your API key: OPENAI_API_KEY=sk-your-key-here\n"
                "3. Or set environment variable: set OPENAI_API_KEY=sk-your-key-here"
            )

        self.client = OpenAI(api_key=self.api_key)
        self.model = model

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract all text content from a .docx file

        Args:
            file_path: Path to the .docx file

        Returns:
            Complete text content with paragraph breaks
        """
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs with style information
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else "Normal"
                # Mark section headers for better structure
                if "Head" in style_name or "Title" in style_name:
                    text_parts.append(f"\n### {text} ###\n")
                else:
                    text_parts.append(text)

        # Also extract tables (some SOL docs use tables)
        for table in doc.tables:
            text_parts.append("\n[TABLE]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
            text_parts.append("[/TABLE]\n")

        return "\n".join(text_parts)

    def extract_structured_data(
        self,
        file_path: str,
        temperature: float = 0.1,
        save_raw_text: bool = False
    ) -> Dict[str, Any]:
        """
        Extract structured data from SOL document using OpenAI

        Args:
            file_path: Path to the .docx file
            temperature: OpenAI temperature parameter (0-1, lower = more deterministic)
            save_raw_text: If True, save extracted text to a .txt file

        Returns:
            Structured JSON data matching the schema
        """
        # Extract text from document
        print(f"Extracting text from {Path(file_path).name}...")
        doc_text = self.extract_text_from_docx(file_path)

        # Optionally save raw text for debugging
        if save_raw_text:
            text_path = Path(file_path).with_suffix('.txt')
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(doc_text)
            print(f"  Saved raw text to {text_path.name}")

        # Call OpenAI API
        print(f"Calling OpenAI API ({self.model})...")
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": OPENAI_EXTRACTION_PROMPT
                    },
                    {
                        "role": "user",
                        "content": f"Extract structured data from this SOL document:\n\n{doc_text}"
                    }
                ],
                temperature=temperature,
                response_format={"type": "json_object"}
            )

            # Parse response
            result = json.loads(response.choices[0].message.content)

            # Add metadata about the extraction
            result["_extraction_metadata"] = {
                "source_file": Path(file_path).name,
                "model": self.model,
                "temperature": temperature,
                "tokens_used": response.usage.total_tokens,
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens
            }

            print(f"  ✓ Extraction complete ({response.usage.total_tokens} tokens)")
            return result

        except Exception as e:
            print(f"  ✗ Error during extraction: {str(e)}")
            raise

    def extract_batch(
        self,
        file_paths: list[str],
        output_dir: str = "sol_formatter/sol_documents",
        save_raw_text: bool = False
    ) -> Dict[str, Any]:
        """
        Extract structured data from multiple documents

        Args:
            file_paths: List of paths to .docx files
            output_dir: Directory to save output files
            save_raw_text: If True, save raw text files

        Returns:
            Summary of batch processing
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        results = []
        successful = 0
        failed = 0
        total_tokens = 0

        print(f"Processing {len(file_paths)} documents...")
        print("=" * 60)

        for idx, file_path in enumerate(file_paths, 1):
            print(f"\n[{idx}/{len(file_paths)}] {Path(file_path).name}")

            try:
                # Extract structured data
                result = self.extract_structured_data(
                    file_path,
                    save_raw_text=save_raw_text
                )

                # Save individual JSON file
                output_filename = Path(file_path).stem + "_structured.json"
                json_path = output_path / output_filename

                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2, ensure_ascii=False)

                print(f"  ✓ Saved to {output_filename}")

                results.append(result)
                successful += 1
                total_tokens += result["_extraction_metadata"]["tokens_used"]

            except Exception as e:
                print(f"  ✗ Failed: {str(e)}")
                failed += 1

        # Save combined results
        combined = {
            "total_documents": len(file_paths),
            "successful": successful,
            "failed": failed,
            "total_tokens_used": total_tokens,
            "documents": results
        }

        combined_path = output_path / "all_structured_documents.json"
        with open(combined_path, 'w', encoding='utf-8') as f:
            json.dump(combined, f, indent=2, ensure_ascii=False)

        # Print summary
        print("\n" + "=" * 60)
        print("BATCH PROCESSING COMPLETE")
        print("=" * 60)
        print(f"Successful: {successful}/{len(file_paths)}")
        print(f"Failed: {failed}/{len(file_paths)}")
        print(f"Total tokens used: {total_tokens:,}")
        print(f"Output directory: {output_path.absolute()}")
        print(f"Combined output: {combined_path.name}")

        return combined

    def validate_output(self, data: Dict[str, Any]) -> bool:
        """
        Validate extracted data against schema

        Args:
            data: Extracted data dictionary

        Returns:
            True if valid, False otherwise
        """
        try:
            # Basic validation - check required fields
            assert "document_metadata" in data
            assert "strands" in data
            assert isinstance(data["strands"], list)

            for strand in data["strands"]:
                assert "strand_code" in strand
                assert "strand_name" in strand
                assert "standards" in strand
                assert isinstance(strand["standards"], list)

                for standard in strand["standards"]:
                    assert "standard_id" in standard
                    assert "standard_statement" in standard
                    assert "knowledge_and_skills" in standard

            return True

        except (AssertionError, KeyError, TypeError) as e:
            print(f"Validation error: {str(e)}")
            return False
